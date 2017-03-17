#!/usr/bin/env python

"""
Create pathology report addendums in Word docx format.

Input: CSV files exported from GA (same as files uploaded to Syapse).

Output: Word doc with Pathogenic and Likeley Pathogenic variants listed,
data from one CSV per page.

"""

import os
import re
import sys
import traceback
from collections import defaultdict
from argparse import ArgumentParser

import docx
import wx
import wx.richtext 

VERSION="1.0"
BUILD="170316"

#----common.py----------------------------------------------------------------

def getScriptPath():
  return os.path.dirname(os.path.realpath(sys.argv[0]))

AAPATT = re.compile('([A-Z][a-z][a-z])')
AA_convert = {
  'Ala': ['A', 'Alanine',], 
  'Arg': ['R', 'Arginine',], 
  'Asn': ['N', 'Asparagine',], 
  'Asp': ['D', 'Aspartic acid',], 
  'Cys': ['C', 'Cysteine',], 
  'Gln': ['Q', 'Glutamine',], 
  'Glu': ['E', 'Glutamic acid',], 
  'Gly': ['G', 'Glycine',], 
  'His': ['H', 'Histidine',], 
  'Ile': ['I', 'Isoleucine',], 
  'Leu': ['L', 'Leucine',], 
  'Lys': ['K', 'Lysine',], 
  'Met': ['M', 'Methionine',], 
  'Phe': ['F', 'Phenylalanine',], 
  'Pro': ['P', 'Proline',], 
  'Pyl': ['O', 'Pyrrolysine',], 
  'Ser': ['S', 'Serine',], 
  'Sec': ['U', 'Selenocysteine',], 
  'Thr': ['T', 'Threonine',], 
  'Trp': ['W', 'Tryptophan',], 
  'Tyr': ['Y', 'Tyrosine',], 
  'Val': ['V', 'Valine',], 
  'Asx': ['B', 'Aspartic acid or Asparagine',], 
  'Glx': ['Z', 'Glutamic acid or Glutamine',], 
  'Xle': ['J', 'Leucine of Isoleucine',], 
  'Ter': ['X', 'Termination codon',], 
}

def aa_replace(matchobj, num=0):
  m = matchobj.group(0)
  expansion = m
  if m in AA_convert:
    expansion = AA_convert[m][num]
  return expansion

def is_float(v):
  try:
    float(v)
  except ValueError:
    return False
  return True

#----classes------------------------------------------------------------------

class GA_CSV:
  required_cols = ['Chr:ChrPos', 'HGVSProtein', 'Gene', 'Pathogenicity',]
  SEVERITY = {
    'Pathogenic': 1,
    'Likely Pathogenic': 2,
  }

  def __init__(self, csvfile):
    self.csvfile = csvfile
    self.fields = []
    self.missing_fields = []
    self.header = []
    self.data = []
    self.datadict = {}
    self.pipeline_version = None
    self.num_variants = len(self.data)
    self.outfile = None
    self.is_valid = False

    self._parse_csv_file()

  def _check_for_required_fields(self, fields):
    self.missing_fields = []
    for col in self.required_cols:
      if not col in fields:
        self.missing_fields.append(col)
    return self.missing_fields

  @staticmethod
  def split_csv_line(line):
    vals = line.rstrip().split(',')
    csv_vals = []
    merge_flag = False
    for v in vals:
      if merge_flag:
        csv_vals[-1] += ",{}".format(v)
      else:
        csv_vals.append(v)
      if v.startswith('"'):
        merge_flag = True
      if v.endswith('"'):
        merge_flag = False
    return csv_vals
  
  @classmethod
  def create_dkey(self, d):
    """Unique key for each variant and allows sorting by pathogenicity 
    severity and gene name"""
    try:
      chrom, chrompos = d['Chr:ChrPos'].split(':')
    except Exception, e:
      sys.stderr.write("  ERROR: {}{}\n\n".format(
                   type(e).__name__, e))
      raise
    severity = self.SEVERITY.get(d['Pathogenicity'], 3)
    gene = d.get('Gene', 'zzzNone')
    if chrom.isdigit():
      chrom = "{:02d}".format(int(chrom))
    try:
      if 'GeneStrand' in d and d['GeneStrand']:
        if d['GeneStrand']=='-':
          chrompos = 300000000-int(chrompos)
    except Exception as e:
      sys.stderr.write("  ERROR: {}{}\n\n".format(type(e).__name__, e))
    dkey = "{} {:<9} {:<2} {:>12} {}".format(severity, gene, chrom, 
       chrompos, d['HGVSProtein'])
  #  sys.stderr.write("chrom %s pos %s pathogenicity %s dkey %s\n" % (chrom, 
  #           chrompos, d['Pathogenicity'], dkey))
    return dkey

  def _parse_csv_file(self):
    fields = None
  #  sys.stderr.write("Parsing: {}\n".format(csvfile))
    with open(self.csvfile, 'r') as fh:
      for line in fh:
        if line.startswith('##'): 
          if line.startswith('##Variants Of'):
            self.is_valid = True
          continue
        if line.startswith('#') and self.required_cols[0] in line:
          fields = self.split_csv_line(line.lstrip('#'))
          self._check_for_required_fields(fields)
          if not self.missing_fields:
            self.fields = fields
        elif line.startswith('#'):
          fields = None
        elif fields and ',' in line:
          vals = self.split_csv_line(line)
          if any(v for v in vals):
            d = dict(zip(fields, vals))
            dkey = self.create_dkey(d)
            self.data.append(d)
            self.datadict[dkey] = d
        elif 'Sample Status Change' in line:
          # pipeline version
          vals = self.split_csv_line(line)
          self.pipeline_version = vals[1]

    self.num_variants = len(self.data)

#-----------------------------------------------------------------------------

def write_progress(boldtext='', normaltext='', bullet=False, newline=True):
  if bullet: 
    sys.stdout.write(' * ')
  sys.stderr.write(boldtext)
  sys.stderr.write(text)
  if newline:
    sys.stderr.write("\n")

def create_word_docx(csvinfo, resident='RESIDENT', signout='ATTENDING', 
    write_progress=write_progress, outfile="pathology_report_addendums.docx"):
  addendum_comment = """This addendum is issued to describe the results of next generation sequencing-based mutational profiling using the Stanford Solid Tumor Actionable Mutation Panel (STAMP), version PIPELINE_VERSION.  All variants considered "pathogenic" or "likely pathogenic" are reported here. For additional details on the variants detected as well as the full list of variants (including variants of uncertain significance) and methodologic details, please see the complete report in EPIC."""
  
  document = docx.Document()
  font = document.styles['Normal'].font
  font.name = 'Times New Roman'
  font.size = docx.shared.Pt(12)
  write_progress("\nCreating report:", " {}".format(outfile))
  for i, (csvfile, info) in enumerate(csvinfo):
    write_progress("  {}) Adding: ".format(i+1), os.path.basename(csvfile))
    if i:
      document.add_page_break()
    p_csvfile = document.add_paragraph()
    p_csvfile.add_run("CSV file: ")
    p_csvfile.add_run(os.path.basename(csvfile))

    p_comment = document.add_paragraph()
    p_comment.add_run('ADDENDUM COMMENT: ').bold = True
    pipeline_version = info.pipeline_version.split('v').pop()
    p_comment.add_run(addendum_comment.replace('PIPELINE_VERSION', pipeline_version))
    
    p_diagnosis = document.add_paragraph()
    run_diagnosis = p_diagnosis.add_run('ADDENDUM DIAGNOSIS:')
    run_diagnosis.bold = True
    run_diagnosis.add_break()
    run_diagnosis2 = p_diagnosis.add_run('SPECIMENID, MUTATIONAL PROFILING BY STAMP')
    run_diagnosis2.bold = True
    run_diagnosis2.add_break()
    have_variants = False
    for vkey, vdata in sorted(info.datadict.items()):
      if vdata['Pathogenicity'] in ('Pathogenic', 'Likely Pathogenic'):
        have_variants = True
        aa_change = AAPATT.sub(aa_replace, vdata['HGVSProtein']).lstrip('p.')
        p_diagnosis.add_run('\t--\tPOSITIVE FOR ').bold=True
        gene_run = p_diagnosis.add_run(vdata['Gene'])
        gene_run.bold=True
        gene_run.italic=True
        run = p_diagnosis.add_run(' {} MUTATION'.format(aa_change))
        run.bold = True
        run.add_break()
    if not have_variants:
      p_diagnosis.add_run('\tNone')
    p_names = document.add_paragraph()
    p_names.add_run('{}/KUNDER/{}'.format(resident, signout)).bold=True
  document.save(outfile)
  write_progress('', "Done.")

def parse_csv_files(csvfiles):
  csvinfo = []
  badfiles = []
  for csvfile in csvfiles:
    ga_csvinfo = GA_CSV(csvfile)
    if ga_csvinfo.missing_fields:
      if ga_csvinfo.is_valid:
        badfiles.append([csvfile, 'Missing required fields: {}'.format(
        ', '.join(ga_csvinfo.missing_fields[:3]))])
      else:
        badfiles.append([csvfile, 'Not a GA CSV'])
    else:
      csvinfo.append([csvfile, ga_csvinfo])
      sys.stdout.write("  {} variants\t{}\n".format(ga_csvinfo.num_variants,
                     os.path.basename(csvfile)))
  return csvinfo, badfiles

def filter_input_files(inputfiles):
  infiles = []
  for in_arg in inputfiles: # input can be files or folders
    if os.path.isfile(in_arg):
      infiles.append(in_arg)
    elif os.path.isdir(in_arg):
      infiles.extend([ os.path.join(in_arg, f) for f in \
               os.listdir(in_arg) ])
  badfiles = []
  csvfiles = []
  for infile in infiles:
    if infile.lower().endswith('.csv'):
      csvfiles.append(infile)
    else:
      badfiles.append(infile)
  return csvfiles, badfiles

#----gui.py-------------------------------------------------------------------

class AddendumApp(wx.App):
  def __init__(self, options, **kwargs):
    self.options = options
    wx.App.__init__(self, kwargs)

  def OnInit(self):
    self.frame = MainFrame(self.options)
    self.frame.Show()
    self.SetTopWindow(self.frame)
    return True

class MainFrame(wx.Frame):
  def __init__(self, options, *args, **kwargs):
    self.options = options
    super(MainFrame, self).__init__(None, *args, size=(550,500),
              title="STAMP Addendum Generator v"+VERSION, **kwargs)
    panel = wx.Panel(self)
    label = wx.StaticText(panel, -1, "Drop STAMP CSV files here:")
    self.rtc = AddendumRTC(panel)
    self.rtc.AddIntroBlurb()
    self.entry_panel = EntryPanel(panel, options.resident, options.signout,
                       options.outfile)
    self.filedrop = FileDropProcessing(self, self.rtc, self.entry_panel)
    self.rtc.SetDropTarget(self.filedrop)

    button_run = wx.Button(panel, -1, "Create report", style=wx.BU_EXACTFIT)
    button_run.SetToolTip(wx.ToolTip("Create Word document"))
    self.Bind(wx.EVT_BUTTON, self.createReport, button_run)
    button_reset = wx.Button(panel, -1, "Reset", style=wx.BU_EXACTFIT)
    button_reset.SetToolTip(wx.ToolTip("Reset CSV list"))
    self.Bind(wx.EVT_BUTTON, self.resetValues, button_reset)
    button_quit = wx.Button(panel, -1, "Quit", style=wx.BU_EXACTFIT)
    button_quit.SetToolTip(wx.ToolTip("Quit application"))
    self.Bind(wx.EVT_BUTTON, self.OnCloseMe, button_quit)
    self.Bind(wx.EVT_CLOSE, self.OnCloseWindow)

    button_sizer = wx.BoxSizer(wx.HORIZONTAL)
    button_sizer.Add(button_run, 0, wx.ALIGN_CENTER_VERTICAL)
    button_sizer.AddStretchSpacer()
    button_sizer.Add(button_reset, 0, wx.ALIGN_CENTER_VERTICAL)
    button_sizer.Add(button_quit, 0, wx.ALIGN_CENTER_VERTICAL)

    sizer = wx.BoxSizer(wx.VERTICAL)
    sizer.Add(label, 0, wx.ALL, border=5)
    sizer.Add(self.rtc, 1, wx.EXPAND|wx.ALL, border=5)
    sizer.Add(self.entry_panel, 0, wx.EXPAND|wx.ALL, border=5)

    sizer.Add(button_sizer, 0, wx.EXPAND|wx.ALL, border=5)
    panel.SetSizer(sizer)


  def createReport(self, event):
    self.filedrop.createReport()

  def resetValues(self, event):
    self.filedrop.reset()

  def OnCloseMe(self, event):
    self.Close(True)

  def OnCloseWindow(self, event):
    self.Destroy()
    
class FileDropProcessing(wx.FileDropTarget):
  def __init__(self, parent, window, entries):
    wx.FileDropTarget.__init__(self)
    self.parent = parent
    self.window = window
    self.entries = entries
    self.current_pos = 0
    self.csvinfo = []
    self.num_csv = 0

  def ScrollWindow(self):
    pos = self.window.GetScrollRange(wx.VERTICAL)
    self.window.Scroll(0, pos)

  def WriteFormattedText(self, boldtext='', normaltext='', bullet=False,
               newline=True):
    self.window.MoveEnd()
    if self.current_pos:
      self.window.SetCaretPosition(self.current_pos)
    if bullet:
      self.window.BeginSymbolBullet('*', 25, 30)
    if boldtext:
      self.window.BeginBold()
      self.window.WriteText(boldtext)
      self.window.EndBold()
    if normaltext:
      self.window.WriteText(normaltext)
    if newline: 
      self.window.Newline()
    if bullet:
      self.window.EndSymbolBullet()
    self.ScrollWindow()
    self.current_pos = self.window.GetCaretPosition()
    self.window.Refresh()

  def OnDropFiles(self, x, y, filenames):
    csvfiles, badfiles = filter_input_files(filenames)
    if badfiles:
      self.window.MoveEnd()
      for badfile in badfiles:
        self.WriteFormattedText(newline=False,
          normaltext="Not a recognized input file: {}\n".\
          format(os.path.basename(badfile)))
        self.ScrollWindow()

    csvinfo, badfiles_csv = parse_csv_files(csvfiles)
    if badfiles_csv:
      self.window.MoveEnd()
      for errmsg, badfile in badfiles_csv:
        self.WriteFormattedText(newline=False,
          normaltext="{}: {}\n".format(errmsg, os.path.basename(badfile)))
        self.ScrollWindow()
    self.WriteFormattedText(newline=True)

    if csvinfo:
      for csvfile, info in csvinfo:
        self.num_csv += 1
        self.WriteFormattedText("CSV file {}: ".format(self.num_csv), os.path.basename(csvfile))
        self.csvinfo.append([csvfile, info])

  def createReport(self):
    try:
      if self.csvinfo:
        self.entries.updateInfo()
        outfile = self.entries.output_file(self.csvinfo[0][0])
        create_word_docx(self.csvinfo, self.entries.resident, 
          self.entries.signout, 
          outfile=outfile, write_progress=self.WriteFormattedText)
        self.WriteFormattedText(newline=True)
    except Exception as e:
      errormsg = "{} {}\n\n".format(type(e).__name__, e)
      self.WriteFormattedText("  ERROR: ", errormsg)
      traceback.print_exc(file=sys.stderr)
      sys.stderr.flush()

  def reset(self):
    self.csvinfo = []
    self.num_csv = 0
    self.WriteFormattedText("CSV files: ", "None")

class AddendumRTC(wx.richtext.RichTextCtrl):
  def __init__(self, parent):
    wx.richtext.RichTextCtrl.__init__(self, parent, -1, "",
            style=wx.TE_READONLY|wx.TE_MULTILINE|wx.HSCROLL)
    self.Bind(wx.EVT_MOUSE_EVENTS, self.DoNothing)

  def DoNothing(self, event):
    pass

  def AddIntroBlurb(self):
    intro_blurb="One Word document will be created for each set of CSV"+\
        " files dropped. The Word document will be saved in the same"+\
        " directory as the first CSV file processed unless otherwise "+\
        "specified."
    intro_items = [   
    ]
    self.BeginFontSize(10)
    self.Newline()
    self.Newline()
    self.WriteText(intro_blurb)
    self.Newline()
    self.BeginSymbolBullet('*', 25, 30)
    for [label, descr] in intro_items:
      self.BeginBold()
      self.WriteText(label)
      self.EndBold()
      self.WriteText(' -- '+descr)
      self.Newline()
    self.EndSymbolBullet()
    self.EndFontSize()
    self.Newline()
    self.Newline()

class EntryPanel(wx.Panel):
  def __init__(self, parent, resident='RESIDENTSNAME',
               signout='SIGNOUTATTENDING', outfile="addendums.docx", **kwargs):
    wx.Panel.__init__(self, parent, **kwargs)
    self.resident = resident
    self.signout = signout
    self.outfile = outfile

    self.createWidgets()
    self.doLayout()
    self.updateInfo(resident, signout, outfile)

  def createWidgets(self):
    self.resident_label = wx.StaticText(self, -1, "Resident's name: ")
    self.resident_entry = wx.TextCtrl(self, -1, self.resident)
    self.signout_label = wx.StaticText(self, -1, "  Sign-out attending: ")
    self.signout_entry = wx.TextCtrl(self, -1, self.signout)
    self.outfile_label = wx.StaticText(self, -1, "  Output file name: ")
    self.outfile_entry = wx.TextCtrl(self, -1, self.outfile)
#    self.resident_entry.Bind(wx.EVT_TEXT, self.updateInfo)
#    self.signout_entry.Bind(wx.EVT_TEXT, self.updateInfo)
#    self.outfile_entry.Bind(wx.EVT_TEXT, self.updateInfo)

  def doLayout(self):
    self.entry_sizer = wx.FlexGridSizer(cols=2, hgap=5, vgap=5)
    self.entry_sizer.AddGrowableCol(1)
    self.entry_sizer.Add(self.resident_label, 0, wx.ALIGN_RIGHT|wx.ALIGN_CENTER_VERTICAL)
    self.entry_sizer.Add(self.resident_entry, 0, wx.EXPAND)
    self.entry_sizer.Add(self.signout_label, 0, wx.ALIGN_RIGHT|wx.ALIGN_CENTER_VERTICAL)
    self.entry_sizer.Add(self.signout_entry, 0, wx.EXPAND)
    self.entry_sizer.Add(self.outfile_label, 0, wx.ALIGN_RIGHT|wx.ALIGN_CENTER_VERTICAL)
    self.entry_sizer.Add(self.outfile_entry, 0, wx.EXPAND)
    self.SetSizer(self.entry_sizer)

  def updateInfo(self, resident='', signout='', outfile=''):
    if resident:
      self.resident_entry.SetValue(resident)
      self.resident=resident
    else:
      self.resident = self.resident_entry.GetValue()
    if signout:
      self.signout_entry.SetValue(signout)
      self.signout=signout
    else:
      self.signout = self.signout_entry.GetValue()
    if outfile:
      self.outfile_entry.SetValue(outfile)
      self.outfile=outfile
    else:
      self.outfile = self.outfile_entry.GetValue()

  def output_file(self, csvreference):
    """Save output in parent directory of csvreference unless output file name
    already specifies valid directory"""
    outfile = None
    if not self.outfile:
      outfile = csvreference.replace('.csv', '') + '_addendums.docx'
    else:
      dirpath = os.path.dirname(self.outfile)
      if os.path.isdir(dirpath):
        outfile = self.outfile
      else:
        csvpath = os.path.dirname(csvreference)
        outfile1 = os.path.abspath(os.path.join(csvpath, os.pardir, self.outfile))
        outfile2 = os.path.abspath(os.path.join(csvpath, os.pardir, os.path.basename(self.outfile)))
        if os.path.isdir(os.path.dirname(outfile1)):
          outfile = outfile1
        else:
          outfile = outfile2
      return outfile


def run_gui(args):
  app = AddendumApp(args)
  app.MainLoop()

#-----------------------------------------------------------------------------
if __name__=='__main__':
  descr = "Create pathology report addendums."
  descr += " All input CSV files will be compiled into one Word document"
  descr += " with data from each CSV file on a separate page."
  parser = ArgumentParser(description=descr)
  parser.add_argument("csv_files", nargs="*",
            help="STAMP CSV files")
  parser.add_argument("-r", "--resident", help="Name of resident",
                      default="RESIDENTSNAME")
  parser.add_argument("-s", "--signout", help="Name of sign-out attending",
                      default="ORIGINALSIGNOUTATTENDING")
  parser.add_argument("-o", "--outfile", help="Name for output file",
                      default="pathology_addendums.docx")
  parser.add_argument("--debug", default=False, action='store_true',
            help="Write debugging messages")

  args = parser.parse_args()
  if len(args.csv_files)==0:
    run_gui(args)
  else:
    csvfiles, badfiles = filter_input_files(args.csv_files)
    sys.stderr.write('\n'.join([ "Not a recognized input file: {}".format(csv) \
                     for csv in badfiles ]))
    csvinfo, badfiles = parse_csv_files(csvfiles)
    sys.stderr.write('\n'.join([ "{}: {}".format(err, csv) for csv, err in badfiles ]))
    create_word_docx(csvinfo, args.resident, args.signout, outfile=args.outfile)


